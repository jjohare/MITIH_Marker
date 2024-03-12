import streamlit as st
import pandas as pd
from pathlib import Path
from docx import Document
from pdf2image import convert_from_path
import pytesseract
import json
from openai import OpenAI
import shutil

def read_text_file(file_input):
    """Read and return the content of a text file or an UploadedFile."""
    if hasattr(file_input, 'read'):  # Checks if it's an UploadedFile or similar
        # It's an UploadedFile; read its content directly
        content = file_input.read()
        # Decode if it's bytes, which is expected for text files
        if isinstance(content, bytes):
            content = content.decode('utf-8')
    else:
        # It's a file path; open and read the file
        file_path = Path(file_input)  # Ensure it's a Path object for consistency
        if file_path.exists():
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
        else:
            content = None  # or raise an error, depending on how you want to handle missing files
    return content

def process_excel_to_csv(uploaded_file, output_dir):
    try:
        # Save temporary file to process
        temp_file_path = output_dir / uploaded_file.name
        with open(temp_file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        workbook = pd.ExcelFile(temp_file_path)
        for sheet_name in workbook.sheet_names:
            df = workbook.parse(sheet_name)
            csv_file_path = output_dir / f"{uploaded_file.name}_{sheet_name}.csv"
            df.to_csv(csv_file_path, index=False, encoding='utf-8')
            st.write(f"Saved {csv_file_path} with UTF-8 encoding.")
    except Exception as e:
        st.error(f"Error processing Excel file {uploaded_file.name}: {e}")

def process_docx_to_txt(uploaded_file, output_dir):
    try:
        temp_file_path = output_dir / uploaded_file.name
        with open(temp_file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        doc = Document(temp_file_path)
        full_text = [para.text for para in doc.paragraphs] + \
                    ['\t'.join(cell.text.strip() for cell in row.cells) for table in doc.tables for row in table.rows]
        output_file_path = output_dir / f"{uploaded_file.name}.txt"
        with open(output_file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_text))
    except Exception as e:
        st.error(f"Error processing DOCX file {uploaded_file.name}: {e}")

def process_pdf_to_txt(uploaded_file, output_dir):
    try:
        temp_file_path = output_dir / uploaded_file.name
        with open(temp_file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        images = convert_from_path(temp_file_path)
        full_text = [pytesseract.image_to_string(image) for image in images]
        output_file_path = output_dir / f"{uploaded_file.name}.txt"
        with open(output_file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_text))
    except Exception as e:
        st.error(f"Error processing PDF file {uploaded_file.name}: {e}")

def concatenate_files(proc_directory):
    texts = []
    for filepath in proc_directory.glob('*.txt'):
        try:
            texts.append(read_text_file(filepath))
        except Exception as e:
            st.error(f"Error reading {filepath}: {e}")
    return "\n\n".join(texts)

def generate_prompt(instructions_text, rubric_text, documents_text):
    """Generate a prompt from instructions, rubric, and concatenated document texts.

    Args:
        instructions_text (str): The instructions text.
        rubric_text (str): The assessment rubric text.
        documents_text (str): The concatenated documents text.

    Returns:
        str: The generated prompt including instructions, rubric, and documents.
    """
    return f"Instructions:\n{instructions_text}\n\nRubric:\n{rubric_text}\n\nDocuments:\n{documents_text}\n\nYOU MUST insert your marks and comments into this section 2 json structure\n{rubric_text}\n\nUse the above structure only."


def call_openai_api(prompt, openai_api_key):
    client = OpenAI(api_key=openai_api_key)
    try:
        model_identifier = "gpt-4-0125-preview"  # Adjust as needed
        response = client.chat.completions.create(
            model=model_identifier,
            messages=[{"role": "user", "content": prompt}]
        )
        # Attempt to parse and return the response as JSON
        try:
            response_json = response.json()
            return response_json
        except ValueError:
            # If response is not JSON, return it as a string
            return {"error": "Received a non-JSON response", "response": str(response)}
    except Exception as e:
        # Handle other exceptions and return the error message
        st.error(f"Error calling OpenAI API: {e}")
        return {"error": str(e)}


def cleanup_directory(directory):
    if directory.exists() and directory.is_dir():
        shutil.rmtree(directory)

def main():
    st.title("Document Processor and OpenAI Analysis")
    api_key = st.text_input("Enter your OpenAI API Key", type="password")
    uploaded_files = st.file_uploader("Choose files", accept_multiple_files=True, type=["pdf", "docx", "xls", "xlsx", "xlsm"])
    instructions_file = st.file_uploader("Upload Instructions (instructions.txt)", type=["txt"])
    rubric_file = st.file_uploader("Upload Rubric (assessment_rubric.json)", type=["json"])

    if st.button("Process and Analyze") and uploaded_files and api_key and instructions_file and rubric_file:
        proc_dir = Path("./temp/processed_files")
        proc_dir.mkdir(parents=True, exist_ok=True)

        instructions_text = read_text_file(instructions_file)
        rubric_text = json.load(rubric_file)

        for uploaded_file in uploaded_files:
            if uploaded_file.name.endswith((".xlsx", ".xls", ".xlsm")):
                process_excel_to_csv(uploaded_file, proc_dir)
            elif uploaded_file.name.endswith(".docx"):
                process_docx_to_txt(uploaded_file, proc_dir)
            elif uploaded_file.name.endswith(".pdf"):
                process_pdf_to_txt(uploaded_file, proc_dir)

        documents_text = concatenate_files(proc_dir)
        prompt = generate_prompt(instructions_text, json.dumps(rubric_text, indent=2), documents_text)
        response = call_openai_api(prompt, api_key)

        if response:
            st.json(response)
        else:
            st.error("Failed to get a response from OpenAI.")

        cleanup_directory(proc_dir)

if __name__ == "__main__":
    main()
